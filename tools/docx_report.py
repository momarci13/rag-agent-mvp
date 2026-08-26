"""DOCX validation report builder.

Builds a draft Word document from a single, domain-agnostic
:class:`agents.risk_schemas.ValidationReport`. This is the source document
for the derived PDF -- see tools/docx_to_pdf.py -- so table/section content
must not drift from tools/pptx_report.py's deck.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from agents.risk_schemas import ValidationReport

DISCLAIMER_BANNER = "DRAFT -- NOT A REGULATORY SUBMISSION"


class ValidationWordReportBuilder:
    """Builds a draft Word validation report from a ValidationReport."""

    def build(self, report: ValidationReport, output_path: str | Path) -> Path:
        doc = Document()

        banner = doc.add_paragraph()
        banner_run = banner.add_run(DISCLAIMER_BANNER)
        banner_run.bold = True
        banner_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_heading(report.title, level=0)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Entity under review: {report.entity_under_review}\n"
            f"Reporting period: {report.reporting_period}\n"
            f"Domain: {report.domain.replace('_', ' ').title()}\n"
            f"Generated: {report.generated_at}\n"
            f"Prepared by: {report.prepared_by}"
        )

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.overall_conclusion or "(pending)")

        doc.add_heading("Scope & Methodology", level=1)
        doc.add_heading("Scope", level=2)
        doc.add_paragraph(report.scope)
        doc.add_heading("Methodology", level=2)
        doc.add_paragraph(report.methodology)

        doc.add_heading("Model / Exposure Overview", level=1)
        overview = doc.add_paragraph()
        overview.add_run(
            f"Entity under review: {report.entity_under_review}\n"
            f"Reporting period: {report.reporting_period}\n"
        )
        for key, value in report.quantitative_results.items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")

        doc.add_heading("Validation Activities Checklist", level=1)
        areas: dict[str, str] = {}
        for f in report.findings:
            if f.area not in areas or f.verdict == "non_compliant":
                areas[f.area] = f.verdict
        if areas:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text = "Area", "Verdict"
            for area, verdict in areas.items():
                row = table.add_row().cells
                row[0].text, row[1].text = area, verdict
        else:
            doc.add_paragraph("No validation areas assessed.")

        doc.add_heading("Findings & Severity Ratings", level=1)
        if report.findings:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, header in enumerate(["Area", "Verdict", "Severity", "Description"]):
                hdr[i].text = header
            for f in report.findings:
                row = table.add_row().cells
                row[0].text = f.area
                row[1].text = f.verdict
                row[2].text = f.severity
                row[3].text = f.description
        else:
            doc.add_paragraph("No findings recorded.")

        doc.add_heading("Quantitative Test Results", level=1)
        if report.quantitative_results:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text = "Metric", "Value"
            for key, value in report.quantitative_results.items():
                row = table.add_row().cells
                row[0].text, row[1].text = str(key), str(value)
        else:
            doc.add_paragraph("No quantitative results were supplied.")

        doc.add_heading("Recommendations & Remediation Plan", level=1)
        actionable = [f for f in report.findings if f.recommendation]
        if actionable:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, header in enumerate(["Recommendation", "Owner", "Deadline (days)"]):
                hdr[i].text = header
            for f in actionable:
                row = table.add_row().cells
                row[0].text = f.recommendation
                row[1].text = f.owner or "Unassigned"
                row[2].text = str(f.remediation_deadline_days) if f.remediation_deadline_days is not None else "n/a"
        else:
            doc.add_paragraph("No outstanding recommendations.")

        doc.add_heading("Overall Validation Conclusion", level=1)
        rating_p = doc.add_paragraph()
        rating_run = rating_p.add_run(f"Overall rating: {report.overall_rating.upper()}")
        rating_run.bold = True
        rating_run.font.size = Pt(14)
        doc.add_paragraph(report.overall_conclusion or "(pending)")

        doc.add_heading("Sign-off", level=1)
        doc.add_paragraph(f"Prepared by: {report.prepared_by}")
        doc.add_paragraph(
            f"Signed off by: {report.signed_off_by or 'PENDING -- human validator sign-off required'}"
        )
        doc.add_paragraph(f"Signed off at: {report.signed_off_at or 'n/a'}")
        disclaimer_p = doc.add_paragraph()
        disclaimer_run = disclaimer_p.add_run(report.disclaimer)
        disclaimer_run.italic = True

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path
