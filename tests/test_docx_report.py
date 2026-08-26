"""Tests for tools/docx_report.py: builds a well-formed report with the
expected content, using python-docx's own read API to inspect it."""

from docx import Document

from agents.risk_schemas import ValidationFinding, ValidationReport
from tools.docx_report import DISCLAIMER_BANNER, ValidationWordReportBuilder


def _sample_report() -> ValidationReport:
    findings = [
        ValidationFinding(
            domain="model_risk", area="Stability testing", verdict="non_compliant",
            severity="critical", description="Population Stability Index exceeds the configured threshold.",
            recommendation="Schedule expedited revalidation.", owner="Model Risk", remediation_deadline_days=14,
        ),
    ]
    return ValidationReport(
        domain="model_risk", title="Model Risk Validation Report -- IFRS9 ECL Retail",
        scope="Scope text", methodology="Methodology text",
        entity_under_review="IFRS9 ECL Retail (M1)", reporting_period="2026Q2",
        findings=findings, quantitative_results={"psi": 0.35},
        overall_rating="unacceptable", overall_conclusion="Model requires urgent revalidation.",
    )


def _full_text(doc: Document) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def test_build_produces_a_well_formed_report(tmp_path):
    report = _sample_report()
    output_path = ValidationWordReportBuilder().build(report, tmp_path / "report.docx")

    assert output_path.exists()
    assert output_path.stat().st_size > 1000

    doc = Document(str(output_path))
    full_text = _full_text(doc)
    assert DISCLAIMER_BANNER in full_text
    assert "Population Stability Index exceeds the configured threshold" in full_text
    assert report.disclaimer in full_text
    assert "PENDING" in full_text


def test_build_reflects_signed_off_state(tmp_path):
    report = _sample_report()
    report.signed_off_by = "john.reviewer"
    report.signed_off_at = "2026-07-01T00:00:00+00:00"
    output_path = ValidationWordReportBuilder().build(report, tmp_path / "signed.docx")
    doc = Document(str(output_path))
    full_text = _full_text(doc)
    assert "john.reviewer" in full_text
    assert "PENDING" not in full_text
